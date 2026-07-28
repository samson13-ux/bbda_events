# -*- coding: utf-8 -*-
"""Génère le mémoire Word BBDA Events (normes Cours_Rédaction scientifique.pptx)."""

from __future__ import annotations

import io
import os
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_BREAK
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn, nsmap
from docx.shared import Cm, Pt, RGBColor, Inches
from PIL import Image

ROOT = Path(__file__).resolve().parents[1]  # memoire/
SHOTS = ROOT / "screenshots"
IMAGES = ROOT / "images"
DATABASE = ROOT / "database"
DIAGRAMMES = ROOT / "docs" / "diagrammes"
OUT = ROOT / "redaction" / "MEMOIRE_BBDA_Events_FOFANA_Samson.docx"

# Largeur max des figures dans le document (cm)
FIG_WIDTH_CM = 14.5


def _set_run_font(run, size=12, bold=False, italic=False, name="Times New Roman"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def _fmt_para(p, size=12, bold=False, italic=False, align="justify", space_after=8, first_line=True):
    p.alignment = {
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
    }.get(align, WD_ALIGN_PARAGRAPH.JUSTIFY)
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(0)
    if first_line and align == "justify":
        pf.first_line_indent = Cm(1.25)
    else:
        pf.first_line_indent = Cm(0)
    for run in p.runs:
        _set_run_font(run, size=size, bold=bold, italic=italic)
    return p


def add_p(doc, text, size=12, bold=False, italic=False, align="justify", space_after=8, first_line=True):
    p = doc.add_paragraph()
    run = p.add_run(text)
    _set_run_font(run, size=size, bold=bold, italic=italic)
    return _fmt_para(p, size=size, bold=bold, italic=italic, align=align, space_after=space_after, first_line=first_line)


def add_title(doc, text, level=1):
    """Titres : 1=chapitre, 2=section, 3=sous-section."""
    sizes = {0: 16, 1: 14, 2: 13, 3: 12}
    p = doc.add_paragraph()
    run = p.add_run(text)
    _set_run_font(run, size=sizes.get(level, 12), bold=True)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT if level else WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_before = Pt(18 if level <= 1 else 12)
    pf.space_after = Pt(10)
    pf.first_line_indent = Cm(0)
    # Bookmark-like outline level via style name if available
    try:
        if level == 0:
            p.style = doc.styles["Title"]
        elif level == 1:
            p.style = doc.styles["Heading 1"]
        elif level == 2:
            p.style = doc.styles["Heading 2"]
        else:
            p.style = doc.styles["Heading 3"]
        # Re-apply Times after style change
        for r in p.runs:
            _set_run_font(r, size=sizes.get(level, 12), bold=True)
    except Exception:
        pass
    return p


def add_page_break(doc):
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def compress_image(path: Path, max_width_px=1400) -> io.BytesIO:
    img = Image.open(path)
    if img.mode in ("RGBA", "P"):
        img = img.convert("RGB")
    w, h = img.size
    if w > max_width_px:
        ratio = max_width_px / float(w)
        img = img.resize((max_width_px, int(h * ratio)), Image.Resampling.LANCZOS)
    buf = io.BytesIO()
    img.save(buf, format="JPEG", quality=82, optimize=True)
    buf.seek(0)
    return buf


def add_figure(doc, filename: str, legend: str, fig_num: int):
    path = SHOTS / filename
    if not path.exists():
        add_p(doc, f"[Figure {fig_num} manquante : {filename}]", italic=True, first_line=False, align="center")
        return
    buf = compress_image(path)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run()
    run.add_picture(buf, width=Cm(FIG_WIDTH_CM))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.first_line_indent = Cm(0)
    cap.paragraph_format.space_after = Pt(12)
    r = cap.add_run(f"Figure {fig_num} — {legend}")
    _set_run_font(r, size=11, italic=True)


def add_photo(doc, path: Path, legend: str, fig_num: int, width_cm=12.0):
    if not path.exists():
        add_p(doc, f"[Photo manquante : {path.name}]", italic=True, first_line=False, align="center")
        return
    buf = compress_image(path, max_width_px=1200)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run()
    run.add_picture(buf, width=Cm(width_cm))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.first_line_indent = Cm(0)
    r = cap.add_run(f"Figure {fig_num} — {legend}")
    _set_run_font(r, size=11, italic=True)


def add_diagramme(doc, filename: str, legend: str, fig_num: int, width_cm=14.0):
    """Insère un diagramme UML depuis memoire/docs/diagrammes/."""
    path = DIAGRAMMES / filename
    if not path.exists():
        add_p(doc, f"[Diagramme manquant : {filename}]", italic=True, first_line=False, align="center")
        return
    add_photo(doc, path, legend, fig_num, width_cm=width_cm)


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        run = p.add_run(h)
        _set_run_font(run, size=11, bold=True)
        p.paragraph_format.first_line_indent = Cm(0)
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = ""
            p = cells[ci].paragraphs[0]
            run = p.add_run(str(val))
            _set_run_font(run, size=10)
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    doc.add_paragraph()
    return table


def add_code_block(doc, code: str, legend: str = ""):
    if legend:
        add_p(doc, legend, italic=True, first_line=False, space_after=4)
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run(code)
    run.font.name = "Courier New"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Courier New")
    run.font.size = Pt(8)


def setup_styles(doc: Document):
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.5)
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    for hname, size in [("Heading 1", 14), ("Heading 2", 13), ("Heading 3", 12)]:
        try:
            hs = doc.styles[hname]
            hs.font.name = "Times New Roman"
            hs.font.size = Pt(size)
            hs.font.bold = True
            hs.font.color.rgb = RGBColor(0, 0, 0)
            hs._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        except Exception:
            pass


def add_footer_note(doc, text):
    add_p(doc, text, size=10, italic=True, first_line=False, align="left", space_after=6)


# ---------------------------------------------------------------------------
# Contenu
# ---------------------------------------------------------------------------

def build():
    doc = Document()
    setup_styles(doc)

    # ========== PAGE DE GARDE ==========
    for _ in range(2):
        add_p(doc, "", first_line=False, space_after=0)
    add_p(doc, "UNIVERSITÉ AUBE NOUVELLE (U-AUBEN)", bold=True, align="center", first_line=False, size=13)
    add_p(doc, "Institut Supérieur d'Information et de Gestion (ISIG International)", align="center", first_line=False, size=12)
    add_p(doc, "Filière : Informatique — Option Programmation", align="center", first_line=False)
    add_p(doc, "", first_line=False)
    add_p(doc, "MÉMOIRE DE FIN D'ÉTUDES", bold=True, align="center", first_line=False, size=14)
    add_p(
        doc,
        "en vue de l'obtention de la Licence en Technologie du Génie Informatique",
        align="center",
        first_line=False,
        italic=True,
    )
    add_p(doc, "", first_line=False)
    add_p(doc, "Thème", bold=True, align="center", first_line=False)
    add_p(
        doc,
        "Conception et développement d'une plateforme web de gestion des déclarations "
        "et de suivi des événements culturels occasionnels au BBDA : de la déclaration "
        "par l'organisateur à la délivrance de la quittance",
        bold=True,
        align="center",
        first_line=False,
        size=12,
    )
    add_p(doc, "", first_line=False)
    add_p(doc, "Présenté et soutenu par :", align="center", first_line=False)
    add_p(doc, "FOFANA Samson", bold=True, align="center", first_line=False, size=13)
    add_p(doc, "", first_line=False)
    add_p(doc, "Structure d'accueil : Bureau Burkinabè du Droit d'Auteur (BBDA)", align="center", first_line=False)
    add_p(doc, "Direction de l'Exploitation, de la Perception et des Affaires Juridiques (DEPAJ)", align="center", first_line=False)
    add_p(doc, "", first_line=False)
    add_p(doc, "Directeur de mémoire : [À COMPLÉTER — nom et grade]", align="center", first_line=False)
    add_p(doc, "Maître de stage : [À COMPLÉTER — nom et fonction]", align="center", first_line=False)
    add_p(doc, "", first_line=False)
    add_p(doc, "Année académique 2025 – 2026", bold=True, align="center", first_line=False)
    add_p(doc, "Ouagadougou, Burkina Faso", align="center", first_line=False)
    add_footer_note(
        doc,
        "Note : remplacer la page de garde par le modèle officiel U-AUBEN si l'établissement l'exige "
        "(cf. Cours_Rédaction scientifique.pptx, diapositive 23).",
    )

    add_page_break(doc)

    # ========== PAGE DE TITRE RÉPÉTÉE ==========
    add_p(doc, "UNIVERSITÉ AUBE NOUVELLE (U-AUBEN)", bold=True, align="center", first_line=False)
    add_p(doc, "ISIG International — Licence Technologie du Génie Informatique", align="center", first_line=False)
    add_p(doc, "", first_line=False)
    add_p(
        doc,
        "Conception et développement d'une plateforme web de gestion des déclarations "
        "et de suivi des événements culturels occasionnels au BBDA",
        bold=True,
        align="center",
        first_line=False,
    )
    add_p(doc, "", first_line=False)
    add_p(doc, "FOFANA Samson", bold=True, align="center", first_line=False)
    add_p(doc, "Année académique 2025 – 2026", align="center", first_line=False)

    add_page_break(doc)

    # ========== DÉDICACE ==========
    add_title(doc, "DÉDICACE", level=0)
    add_p(
        doc,
        "[À COMPLÉTER PAR L'ÉTUDIANT — texte personnel, une demi-page à une page. "
        "Proposition de départ à personnaliser :] ",
        italic=True,
        first_line=False,
    )
    add_p(
        doc,
        "À ma famille, pour son soutien constant tout au long de mon parcours universitaire. "
        "À mes formateurs de l'Université Aube Nouvelle, qui m'ont transmis les bases "
        "méthodologiques et techniques nécessaires à la réalisation de ce travail. "
        "Au personnel du Bureau Burkinabè du Droit d'Auteur, pour l'accueil et la confiance "
        "accordés durant ce stage.",
        align="center",
        first_line=False,
        italic=True,
    )

    add_page_break(doc)

    # ========== REMERCIEMENTS ==========
    add_title(doc, "REMERCIEMENTS", level=0)
    add_p(
        doc,
        "[À PERSONNALISER — noms exacts du directeur de mémoire, du maître de stage et des agents DEPAJ.]",
        italic=True,
        first_line=False,
    )
    add_p(
        doc,
        "Au terme de ce travail, nous tenons à exprimer notre profonde gratitude à toutes "
        "les personnes qui ont contribué, de près ou de loin, à sa réalisation.",
    )
    add_p(
        doc,
        "Nos remerciements s'adressent en premier lieu à la Direction du Bureau Burkinabè "
        "du Droit d'Auteur (BBDA) pour nous avoir accueilli au sein de la structure, ainsi "
        "qu'à notre maître de stage [À COMPLÉTER] pour son encadrement, sa disponibilité "
        "et les précisions métier apportées tout au long du projet. Nous remercions également "
        "les agents de la Direction de l'Exploitation, de la Perception et des Affaires "
        "Juridiques (DEPAJ) pour le temps consacré aux entretiens et à l'explication du "
        "circuit papier des séances occasionnelles.",
    )
    add_p(
        doc,
        "Nous remercions notre directeur de mémoire [À COMPLÉTER], pour ses conseils "
        "méthodologiques et son suivi académique. Nos remerciements vont aussi à l'ensemble "
        "du corps enseignant de l'ISIG International / Université Aube Nouvelle, ainsi qu'à "
        "nos proches pour leur soutien moral.",
    )

    add_page_break(doc)

    # ========== RÉSUMÉ / MOTS-CLÉS ==========
    add_title(doc, "RÉSUMÉ", level=0)
    add_p(
        doc,
        "Le présent mémoire porte sur la conception et le développement d'une plateforme web "
        "baptisée BBDA Events, destinée à dématérialiser le circuit de déclaration, "
        "d'évaluation, de paiement et de délivrance de quittance des événements culturels "
        "occasionnels au Bureau Burkinabè du Droit d'Auteur (BBDA). Contrairement aux "
        "utilisateurs permanents déjà couverts par une application existante, les "
        "organisateurs de séances occasionnelles (concerts, festivals, galas, spectacles) "
        "relevaient encore d'un processus entièrement manuel fondé sur des fiches papier. "
        "Après analyse du processus à la DEPAJ et formalisation des règles métier, une "
        "application Flask (architecture MVC) a été réalisée : inscription et déclaration "
        "en ligne, espace agent (fixation manuelle du tarif et de la redevance, confirmation "
        "de paiement), génération automatique de quittance PDF, notifications par courriel, "
        "gestion des arriérés avec seuil de blocage paramétrable, espace administrateur et "
        "face publique de promotion des événements quittancés. La solution est déployée en "
        "ligne (Render) et validée par une suite de tests automatisés (plus de cent tests) "
        "ainsi qu'un parcours d'intégration bout-en-bout. Les hypothèses de réduction des "
        "délais et d'amélioration du recouvrement sont discutées au regard des mécanismes "
        "implémentés. Des perspectives (paiement en ligne, billetterie, diagrammes UML "
        "formels) sont proposées.",
        first_line=False,
    )
    add_p(doc, "Mots-clés : droit d'auteur ; gestion collective ; BBDA ; séance occasionnelle ; "
          "dématérialisation ; Flask ; quittance PDF ; arriérés ; Burkina Faso.",
          italic=True, first_line=False)

    add_p(doc, "ABSTRACT", bold=True, align="center", first_line=False, size=13)
    add_p(
        doc,
        "This bachelor's thesis presents the design and development of BBDA Events, a web "
        "platform that digitizes the declaration, assessment, payment and receipt workflow "
        "for occasional cultural events at the Burkinabe Copyright Office (BBDA). Occasional "
        "event organizers previously relied on a fully paper-based process, unlike permanent "
        "users already covered by an existing information system. Using a Flask MVC "
        "architecture, the platform delivers online declarations, agent-side fee setting, "
        "PDF receipt generation, email notifications, arrears management with account "
        "blocking, an admin console and a public promotion face for cleared events. The "
        "system is deployed online and validated through automated tests. Research "
        "hypotheses on processing time and recovery of unpaid fees are discussed against "
        "the implemented features.",
        first_line=False,
        italic=True,
    )
    add_p(doc, "Keywords: copyright; collective management; BBDA; occasional session; "
          "digitization; Flask; PDF receipt; arrears; Burkina Faso.",
          italic=True, first_line=False)

    add_page_break(doc)

    # ========== SOMMAIRE ==========
    add_title(doc, "SOMMAIRE", level=0)
    sommaire = [
        "Dédicace",
        "Remerciements",
        "Résumé / Abstract",
        "Liste des figures",
        "Liste des tableaux",
        "Introduction générale",
        "Chapitre 1 — Synthèse bibliographique et étude de l'existant",
        "    1.1 Définitions des concepts clés",
        "    1.2 Approches théorique et empirique (revue de littérature)",
        "    1.3 Cadre juridique et institutionnel burkinabè",
        "    1.4 État des lieux au BBDA et expériences comparées",
        "Chapitre 2 — Matériel et méthodes",
        "    2.1 Présentation de la structure d'accueil",
        "    2.2 Démarche méthodologique",
        "    2.3 Conception UML (cas d'utilisation, classes, activité, déploiement)",
        "    2.4 Outils et choix techniques",
        "    2.5 Modélisation des données et règles de gestion",
        "Chapitre 3 — Présentation et analyse des résultats",
        "    3.1 à 3.18 — Modules livrés (authentification, déclaration, agent, quittance, etc.)",
        "    3.19 Discussion, critiques et suggestions",
        "    3.20 Bilan du stage",
        "Conclusion générale",
        "Références bibliographiques",
        "Annexes",
        "Table des matières détaillée",
        "Annexe finale — Points à compléter / vérifier par l'étudiant",
    ]
    for line in sommaire:
        add_p(doc, line, first_line=False, space_after=2, align="left")

    add_page_break(doc)

    # ========== LISTE DES FIGURES ==========
    add_title(doc, "LISTE DES FIGURES", level=0)
    figures = [
        (1, "01-accueil.png / 07-accueil-v2.png", "Page d'accueil publique de BBDA Events"),
        (2, "02-inscription.png", "Formulaire d'inscription d'un organisateur"),
        (3, "03-connexion.png", "Formulaire de connexion"),
        (4, "08-flash-erreur.png", "Message d'erreur (compte inconnu)"),
        (5, "09-dashboard-orga1.png", "Tableau de bord organisateur"),
        (6, "10-dashboard-orga4-bloque.png", "Compte organisateur bloqué pour arriéré"),
        (7, "11-formulaire-concert.png", "Formulaire de déclaration — concert"),
        (8, "12-formulaire-festival-artistes.png", "Formulaire — festival avec plusieurs artistes"),
        (9, "13-formulaire-erreurs.png", "Validation des erreurs de saisie"),
        (10, "15-detail-orga3-0.png", "Détail d'une déclaration (vue d'ensemble)"),
        (11, "15-detail-orga3-1.png", "Chronologie du traitement"),
        (12, "16-detail-montant-fixe.png", "Détail après fixation du montant"),
        (13, "17-dashboard-agent.png", "Tableau de bord de l'agent BBDA"),
        (14, "18-agent-surveillance.png", "Comptes sous surveillance"),
        (15, "19-agent-arrieres.png", "Organisateurs en situation d'arriéré"),
        (16, "20-agent-traitement.png", "Traitement d'une déclaration (agent)"),
        (17, "21-avant-validation-total.png", "Calcul temps réel Tarif + Redevance"),
        (18, "24-formulaire-paiement-vide.png", "Formulaire de confirmation de paiement"),
        (19, "25-formulaire-paiement-rempli.png", "Paiement rempli (exemple chèque)"),
        (20, "27-dossier-apres-paiement.png", "Dossier après paiement / quittance"),
        (21, "30-quittance-pdf-v3.png", "Quittance PDF générée automatiquement"),
        (22, "31-quittance-partiel.png", "Quittance / mention paiement partiel"),
        (23, "32-detail-avec-bouton-telecharger.png", "Téléchargement de la quittance"),
        (24, "33-email-montant-fixe.png", "Email automatique — montant fixé"),
        (25, "34-detail-evenement-public.png", "Page publique d'un événement promu"),
        (26, "35-formulaire-promotion.png", "Section promotion du formulaire"),
        (27, "06-espace-admin.png", "Espace administrateur"),
        (28, "images/*.jpg", "Quittance papier BBDA (référence terrain)"),
        (29, "database/schema_diagram.pdf", "Diagramme relationnel (annexe)"),
        (30, "docs/diagrammes/01-cas-utilisation.png", "Diagramme de cas d'utilisation"),
        (31, "docs/diagrammes/02-classes.png", "Diagramme de classes"),
        (32, "docs/diagrammes/05-activite.png", "Diagramme d'activité du circuit métier"),
        (33, "docs/diagrammes/06-deploiement.png", "Diagramme de déploiement"),
        (34, "docs/diagrammes/03-sequence-declaration.png", "Diagramme de séquence — déclaration"),
        (35, "docs/diagrammes/04-sequence-paiement-quittance.png", "Diagramme de séquence — paiement / quittance"),
    ]
    for num, fich, leg in figures:
        add_p(doc, f"Figure {num} — {leg}  ({fich})", first_line=False, space_after=2, align="left", size=11)

    add_title(doc, "LISTE DES TABLEAUX", level=2)
    for t in [
        "Tableau 1 — Comparaison utilisateurs permanents / organisateurs occasionnels",
        "Tableau 2 — Stack technique retenue et justifications",
        "Tableau 3 — Principales règles métier (extrait)",
        "Tableau 4 — Cycle des statuts d'une déclaration",
        "Tableau 5 — Objectifs spécifiques et livrables correspondants",
    ]:
        add_p(doc, t, first_line=False, space_after=2, align="left", size=11)

    add_page_break(doc)

    # ========== INTRODUCTION ==========
    add_title(doc, "INTRODUCTION GÉNÉRALE", level=0)

    add_title(doc, "Généralités sur le thème", level=2)
    add_p(
        doc,
        "La gestion collective des droits d'auteur consiste, pour une organisation habilitée, "
        "à percevoir pour le compte des créateurs les redevances dues par toute personne "
        "exploitant leurs œuvres, puis à leur en reverser le produit. Au Burkina Faso, cette "
        "mission est assurée par le Bureau Burkinabè du Droit d'Auteur (BBDA), établissement "
        "public à caractère professionnel créé en 1985, chargé de la protection et de la "
        "gestion collective des droits d'auteur et des droits voisins (BBDA, s.d. ; "
        "Wikipedia, 2024). Le BBDA distingue deux grandes catégories d'utilisateurs "
        "d'œuvres protégées : les utilisateurs permanents (radios, télévisions, débits de "
        "boissons, hôtels, etc.), qui exploitent des œuvres de manière continue et doivent "
        "obtenir une autorisation préalable, et les organisateurs d'événements culturels "
        "occasionnels (concerts, festivals, galas, spectacles, foires…), relevant d'un "
        "circuit de déclaration / perception propre aux « séances occasionnelles » "
        "(Sawadogo, 2020 ; Kulture Kibaré, 2022).",
    )
    add_p(
        doc,
        "À ce jour, le BBDA dispose déjà d'un outil numérique dédié aux utilisateurs "
        "permanents. En revanche, le circuit de déclaration des événements occasionnels "
        "restait, au démarrage du stage, entièrement manuel et fondé sur des supports "
        "papier, de la fiche de déclaration remplie par l'organisateur jusqu'à la "
        "quittance délivrée après paiement (Protocole de stage BBDA / U-AUBEN, 2025-2026). "
        "C'est cette absence de couverture numérique qui a motivé le présent travail et "
        "conduit à la conception de la plateforme BBDA Events, accessible en ligne à "
        "l'adresse https://bbda-events.onrender.com.",
    )

    add_title(doc, "Justification du choix du sujet et motivations", level=2)
    add_p(
        doc,
        "Le choix du sujet répond à un besoin concret exprimé par la structure d'accueil : "
        "moderniser un processus administratif encore intégralement papier, source de délais "
        "de traitement allongés et d'une charge importante pour les agents de la Direction "
        "de l'Exploitation, de la Perception et des Affaires Juridiques (DEPAJ). Sur le plan "
        "académique, le sujet mobilise la conception de systèmes d'information (modélisation "
        "de données, architecture logicielle), le développement web (back-end et front-end) "
        "et la gestion de projet (recueil de besoins, arbitrages fonctionnels, tests), sur un "
        "cas réel à fort enjeu : le recouvrement de redevances qui bénéficient in fine aux "
        "créateurs burkinabè. La pertinence du thème s'inscrit également dans la dynamique "
        "de digitalisation engagée par le BBDA en 2025 (plateforme d'adhésion / déclaration "
        "d'œuvres et application B-GEOLOC), qui confirme l'orientation institutionnelle vers "
        "la numérisation des services (leFaso.net, 2025 ; Digital Magazine Burkina, 2025).",
    )

    add_title(doc, "Identification et formulation du problème", level=2)
    add_p(
        doc,
        "Le processus actuellement décrit pour les événements occasionnels se déroule ainsi : "
        "l'organisateur se déplace au siège du BBDA, remplit une fiche de déclaration papier ; "
        "un agent de la DEPAJ analyse les informations et fixe manuellement, sur une fiche "
        "d'évaluation, le montant dû (tarif de référence et redevance complémentaire) ; après "
        "paiement, une quittance physique est délivrée. Ce circuit présente plusieurs limites : "
        "délais allongés avec le volume de dossiers, charge de ressaisie et de classement "
        "papier, et surtout absence de suivi centralisé des paiements et des arriérés "
        "(redevances restées impayées en tout ou en partie).",
    )
    add_p(
        doc,
        "Le problème peut se formuler ainsi : comment le BBDA peut-il fiabiliser, accélérer "
        "et rendre traçable le traitement des déclarations d'événements culturels occasionnels "
        "et le recouvrement des redevances associées, alors que ce processus repose aujourd'hui "
        "intégralement sur des supports papier et des opérations manuelles ?",
    )

    add_title(doc, "Questions de recherche", level=2)
    add_p(
        doc,
        "Question générale : dans quelle mesure la dématérialisation du processus de "
        "déclaration et de traitement des événements occasionnels peut-elle améliorer "
        "l'efficacité administrative du BBDA et la traçabilité des paiements ?",
        first_line=False,
    )
    add_p(doc, "Questions spécifiques :", first_line=False)
    add_p(doc, "1. Quelles sont les étapes et les règles de gestion exactes du processus actuel "
          "de déclaration, d'évaluation et de paiement au BBDA ?", first_line=False)
    add_p(doc, "2. Quelle architecture logicielle et quel modèle de données permettent de "
          "couvrir fidèlement ce processus tout en restant évolutifs ?", first_line=False)
    add_p(doc, "3. Une plateforme web permet-elle de réduire les délais de traitement et de "
          "sécuriser le recouvrement des redevances, notamment via un mécanisme de suivi "
          "des arriérés et de notifications automatiques ?", first_line=False)

    add_title(doc, "Énoncé des objectifs de recherche", level=2)
    add_p(
        doc,
        "Objectif général. Concevoir et développer une plateforme web de gestion et de suivi "
        "des déclarations d'événements culturels occasionnels au BBDA, couvrant l'ensemble "
        "du processus depuis la déclaration de l'organisateur jusqu'à la délivrance de la "
        "quittance, en intégrant le suivi des paiements, la gestion des arriérés et un "
        "système de notifications automatiques.",
    )
    add_p(doc, "Objectifs spécifiques :", first_line=False)
    objs = [
        "Analyser le processus actuel de déclaration, d'évaluation et de paiement des redevances au BBDA.",
        "Identifier les besoins fonctionnels et techniques des agents de la DEPAJ et des organisateurs.",
        "Concevoir l'architecture fonctionnelle et technique (modèle de données, règles métier, architecture logicielle).",
        "Développer un module de déclaration en ligne pour les organisateurs.",
        "Développer un espace agent permettant l'évaluation manuelle du montant dû et le suivi des paiements.",
        "Développer un système de notifications automatiques par courrier électronique.",
        "Développer un module de gestion des arriérés (alertes, surveillance, blocage).",
        "Développer un module de génération automatique de la quittance au format PDF.",
        "Tester, valider et documenter l'ensemble de la plateforme développée.",
    ]
    for i, o in enumerate(objs, 1):
        add_p(doc, f"{i}. {o}", first_line=False, space_after=4)

    add_title(doc, "Formulation des hypothèses", level=2)
    add_p(
        doc,
        "H1. La dématérialisation du processus de déclaration — fondée sur une déclaration "
        "de l'organisateur plutôt que sur une demande d'autorisation préalable du type "
        "utilisateurs permanents — permet de réduire significativement les délais de "
        "traitement et la charge administrative des agents, tout en améliorant la "
        "traçabilité des déclarations et des paiements.",
    )
    add_p(
        doc,
        "H2. La mise en place d'un système centralisé de gestion des arriérés, assorti "
        "d'alertes automatiques et d'un mécanisme de blocage des comptes débiteurs, "
        "facilite le recouvrement des redevances impayées et réduit les pertes "
        "financières pour le BBDA.",
    )

    add_title(doc, "Annonce du plan", level=2)
    add_p(
        doc,
        "Le présent mémoire s'articule en trois chapitres, conformément à la structure "
        "recommandée par le cours d'Initiation à la méthodologie de recherche de "
        "l'Université Aube Nouvelle (janvier 2026). Le premier chapitre constitue la "
        "synthèse bibliographique et l'étude de l'existant : définitions, revue de "
        "littérature, cadre juridique et état des lieux. Le deuxième chapitre présente "
        "le matériel et les méthodes : structure d'accueil, démarche, choix techniques "
        "et modélisation. Le troisième chapitre présente et analyse les résultats "
        "module par module, avant discussion critique et bilan du stage. Le mémoire se "
        "conclut par un bilan général, la confrontation des hypothèses aux résultats, "
        "les limites et les perspectives.",
    )

    add_page_break(doc)

    # ========== CHAPITRE 1 ==========
    add_title(doc, "CHAPITRE 1 — SYNTHÈSE BIBLIOGRAPHIQUE ET ÉTUDE DE L'EXISTANT", level=1)

    add_title(doc, "1.1 Définitions des concepts clés", level=2)
    defs = [
        ("Bureau Burkinabè du Droit d'Auteur (BBDA)",
         "Institution publique burkinabè chargée de la protection et de la gestion "
         "collective des droits d'auteur et des droits voisins ; elle perçoit, pour le "
         "compte des ayants droit, les redevances dues par les exploitants d'œuvres protégées."),
        ("Direction de l'Exploitation, de la Perception et des Affaires Juridiques (DEPAJ)",
         "Direction du BBDA en charge, notamment, du traitement des déclarations "
         "d'événements culturels occasionnels."),
        ("Séance / événement culturel occasionnel",
         "Manifestation ponctuelle (concert, festival, spectacle, gala, foire, défilé…) "
         "donnant lieu à l'exploitation publique d'œuvres protégées, par opposition à "
         "l'exploitation permanente."),
        ("Déclaration",
         "Acte par lequel un organisateur signale au BBDA la tenue prochaine d'un "
         "événement occasionnel. Dans le cadrage du protocole de stage, ce processus "
         "se distingue de la demande d'autorisation préalable applicable aux utilisateurs permanents."),
        ("Tarif et redevance",
         "Le tarif est le montant de référence tiré du barème interne du BBDA selon le "
         "type d'événement ; la redevance est un montant complémentaire apprécié par "
         "l'agent selon le contexte (jauge, notoriété, prix des billets…). La somme "
         "des deux constitue le montant total dû."),
        ("Quittance",
         "Document attestant du paiement de la redevance ; historiquement remis en main "
         "propre sur papier, produit au format PDF par la plateforme."),
        ("Arriéré",
         "Partie de la redevance restée impayée après un paiement partiel, ou totalité "
         "en l'absence de paiement."),
        ("Compte sous surveillance",
         "Statut appliqué à un compte organisateur suspecté de fraude ou injoignable, "
         "destiné à déclencher une alerte à sa prochaine connexion."),
    ]
    for titre, texte in defs:
        add_p(doc, f"{titre}. {texte}")

    add_title(doc, "1.2 Approches théorique et empirique (revue de littérature)", level=2)
    add_p(
        doc,
        "La revue de littérature mobilise trois champs : (i) la gestion collective du "
        "droit d'auteur et son cadre légal burkinabè ; (ii) la dématérialisation des "
        "procédures administratives et des systèmes d'information de gestion ; "
        "(iii) les expériences d'organismes de gestion collective (OGC) comparables "
        "en Afrique et en Europe. Les sources sont présentées par auteur / année "
        "(ou organisme / année) conformément au cours de méthodologie.",
    )
    add_p(
        doc,
        "Sur le plan juridique, la Loi n° 048-2019/AN du 12 novembre 2019 portant "
        "protection de la propriété littéraire et artistique (promulguée par le Décret "
        "n° 2019-1275/PRES du 23 décembre 2019) constitue le texte de référence au "
        "Burkina Faso ; elle abroge la Loi n° 032-99/AN du 22 décembre 1999 "
        "(WIPO Lex, 2019). Le titre relatif à la gestion collective confie à "
        "l'organisme professionnel de gestion collective (en pratique le BBDA) la "
        "mission de gérer les droits et de concéder licences et autorisations pour le "
        "compte des titulaires (art. 102 et 103, tels que présentés dans les analyses "
        "de la presse culturelle et dans le dossier WIPO Lex). L'Arrêté "
        "n° 01-053/MAC/SG/BBDA du 20 mars 2000 précise historiquement les règles de "
        "perception des redevances auprès des utilisateurs et organisateurs de "
        "spectacles (WIPO Lex, 2000). Un guide pratique BBDA / OMPI (éd. révisée 2011) "
        "accompagne les utilisateurs d'œuvres protégées.",
    )
    add_p(
        doc,
        "Sur le plan empirique local, Sawadogo (2020) décrit le processus de collecte "
        "des droits de la séance occasionnelle au BBDA (autorisation / déclaration "
        "préalable, redevance, relevé d'exploitation), en s'appuyant sur des propos "
        "du service compétent. Kulture Kibaré (2022) vulgarise le fonctionnement des "
        "séances occasionnelles auprès du public culturel. Les médias généralistes "
        "(Burkina24, 2022 ; Sidwaya, 2024-2025 ; Actualité.bf, 2025) documentent "
        "périodiquement les sessions de répartition de droits (centaines de millions "
        "FCFA redistribués à des milliers de bénéficiaires), illustrant l'enjeu "
        "économique de la collecte, dont une part est liée à la représentation "
        "directe / séances occasionnelles.",
    )
    add_p(
        doc,
        "Sur le plan comparatif, la SACEM (France) propose une autorisation et une "
        "déclaration d'événements largement dématérialisées (SACEM, s.d.). Au Sénégal, "
        "la SODAV distingue contrat de représentation générale et séance occasionnelle, "
        "avec autorisation préalable, paiement de redevance, quittance et programmes "
        "d'œuvres pour la répartition (SODAV, s.d.). En Afrique du Sud, SAMRO délivre "
        "des licences d'usage public de musique via un portail Music User (SAMRO, s.d.). "
        "Ces expériences convergent vers un schéma déclaration / licence → paiement → "
        "preuve (quittance) → documentation des œuvres, que BBDA Events transpose au "
        "contexte burkinabè pour le seul périmètre organisateur occasionnel.",
    )
    add_p(
        doc,
        "Évaluation critique. Les sources juridiques (WIPO Lex) offrent une base "
        "solide mais doivent être citées après vérification du PDF officiel. La presse "
        "culturelle burkinabè est précieuse pour décrire la pratique métier, mais "
        "n'équivaut pas à une étude académique peerrreviewed ; elle est utilisée ici "
        "comme matériau empirique complémentaire aux documents internes du stage "
        "(protocole, cahier des charges). Les comparaisons internationales éclairent "
        "l'état de l'art fonctionnel sans prétendre à une transposition juridique "
        "directe. [À APPROFONDIR : articles académiques CISAC / OMPI sur la gestion "
        "collective en Afrique — voir annexe « Points à compléter ».]",
        italic=False,
    )

    add_title(doc, "1.3 Cadre juridique et institutionnel burkinabè", level=2)
    add_p(
        doc,
        "Créé en 1985, le BBDA assure la gestion collective au plan national et gère "
        "également, via des accords de réciprocité, les intérêts d'organismes étrangers "
        "(BBDA, s.d.). La distinction utilisateurs permanents / utilisateurs occasionnels "
        "structure l'organisation des services et justifie des outils informatiques "
        "distincts. Le Tableau 1 résume cette distinction, telle que cadrée par le "
        "protocole de stage.",
    )
    add_p(doc, "Tableau 1 — Comparaison utilisateurs permanents / organisateurs occasionnels",
          bold=True, first_line=False, align="center", size=11)
    add_table(
        doc,
        ["Critère", "Utilisateurs permanents", "Organisateurs occasionnels (périmètre du stage)"],
        [
            ["Processus", "Demande d'autorisation préalable", "Déclaration de l'événement"],
            ["Fréquence", "Exploitation continue", "Manifestation ponctuelle"],
            ["Outil existant avant le stage", "Application dédiée BBDA", "Aucun outil numérique (papier)"],
            ["Preuve de paiement", "Selon procédures internes", "Quittance"],
        ],
    )
    add_p(
        doc,
        "Précision méthodologique importante. La presse culturelle évoque souvent une "
        "« autorisation » à solliciter environ 72 heures avant une séance occasionnelle "
        "(Sawadogo, 2020). Le protocole de stage, validé comme cadrage académique du "
        "projet, insiste sur la distinction déclaration (occasionnels) versus "
        "autorisation préalable (permanents). Dans la plateforme BBDA Events, le "
        "vocabulaire métier retenu est celui de la déclaration, de l'évaluation et de "
        "la quittance, conformément au cahier des charges. [À VÉRIFIER AVEC LE MAÎTRE "
        "DE STAGE : formulation officielle exacte utilisée à la DEPAJ.]",
    )

    add_title(doc, "1.4 État des lieux au BBDA et expériences comparées", level=2)
    add_p(
        doc,
        "L'analyse de l'existant s'appuie sur trois documents de cadrage : le protocole "
        "de stage (2025-2026), le cahier des charges V2 BBDA Events et le guide de "
        "développement séquencé en vingt étapes. Ces documents confirment l'absence "
        "d'outil numérique pour les séances occasionnelles et listent les modules "
        "attendus (déclaration, agent, notifications email, arriérés, quittance PDF, "
        "tests, mémoire). Des incohérences mineures entre documents (périmètre "
        "WhatsApp reporté faute d'API payante, calcul automatique de redevance écarté "
        "au profit d'une fixation manuelle RM-030) ont été arbitrées au profit des "
        "règles métier formalisées dans REGLES_METIER.md.",
    )
    add_p(
        doc,
        "En avril 2025, le BBDA a lancé une plateforme numérique d'adhésion et de "
        "déclaration d'œuvres ainsi que l'application mobile B-GEOLOC pour la collecte "
        "terrain (leFaso.net, 2025 ; Digital Magazine Burkina, 2025). BBDA Events ne "
        "duplique pas ces outils : il couvre le circuit organisateur → quittance, encore "
        "peu digitalisé, et s'inscrit dans la même trajectoire de modernisation.",
    )

    add_page_break(doc)

    # ========== CHAPITRE 2 ==========
    add_title(doc, "CHAPITRE 2 — MATÉRIEL ET MÉTHODES", level=1)

    add_title(doc, "2.1 Présentation de la structure d'accueil et du cadre de l'étude", level=2)
    add_p(
        doc,
        "Le stage d'une durée de trois mois (année académique 2025-2026) s'est déroulé "
        "au Bureau Burkinabè du Droit d'Auteur (BBDA), à Ouagadougou, au contact de la "
        "DEPAJ. [À COMPLÉTER : dates exactes de début et de fin de stage, nom du service "
        "d'affectation quotidien, organigramme simplifié fourni par le BBDA si disponible.]",
    )
    add_p(
        doc,
        "Le BBDA a pour missions principales la protection des droits, la perception des "
        "redevances et la répartition aux ayants droit. L'adresse institutionnelle "
        "couramment communiquée est : 01 BP 3926 Ouagadougou 01 — Tél. : 25 32 47 50 "
        "(reprise également sur les quittances générées).",
    )

    add_title(doc, "2.2 Démarche méthodologique", level=2)
    add_p(
        doc,
        "La démarche retenue est de type mixte, à dominante qualitative et "
        "expérimentale (réalisation d'un artefact logiciel), conforme à une recherche "
        "appliquée en génie informatique :",
    )
    add_p(doc, "Phase 1 — Analyse. Observation du processus à la DEPAJ, entretiens avec les "
          "agents, collecte des fiches papier (déclaration, évaluation) et photographie "
          "d'une quittance physique de référence (dossier images/).", first_line=False)
    add_p(doc, "Phase 2 — Conception. Modélisation UML (cas d'utilisation, classes, "
          "séquence, activité, déploiement), schéma relationnel, règles RM-XXX, "
          "architecture MVC Flask, spécification des rôles (organisateur, agent, "
          "administrateur).", first_line=False)
    add_p(doc, "Phase 3 — Réalisation. Développement itératif en vingt « prompts » "
          "(guide de développement V2), avec tests automatisés à chaque étape.", first_line=False)
    add_p(doc, "Phase 4 — Validation et documentation. Tests fonctionnels et d'intégration "
          "(pytest, Playwright), données de démonstration, scénario de soutenance, "
          "rédaction du mémoire.", first_line=False)

    add_title(doc, "2.3 Conception UML", level=2)
    add_p(
        doc,
        "Conformément au protocole de stage, la conception a été formalisée par un "
        "ensemble de diagrammes UML. Les sources Mermaid et les exports PNG sont "
        "conservés dans le dossier memoire/docs/diagrammes/ pour consultation et "
        "réédition. Les figures ci-dessous présentent les vues principales ; les "
        "diagrammes de séquence détaillés sont repris en annexe.",
    )
    add_p(
        doc,
        "Le diagramme de cas d'utilisation identifie les acteurs (organisateur, agent "
        "BBDA, administrateur, public) et les principales fonctionnalités du système.",
    )
    add_diagramme(
        doc,
        "01-cas-utilisation.png",
        "Diagramme de cas d'utilisation de la plateforme BBDA Events",
        30,
        width_cm=13.5,
    )
    add_p(
        doc,
        "Le diagramme de classes reprend les entités du modèle de données "
        "(utilisateur, organisateur, déclaration, évaluation, paiement, quittance, "
        "arriéré, etc.) et leurs cardinalités.",
    )
    add_diagramme(
        doc,
        "02-classes.png",
        "Diagramme de classes (extrait du modèle SQLAlchemy)",
        31,
        width_cm=13.5,
    )
    add_p(
        doc,
        "Le diagramme d'activité décrit le circuit métier complet, de l'inscription "
        "à la délivrance de la quittance et à l'éventuelle publication publique.",
    )
    add_diagramme(
        doc,
        "05-activite.png",
        "Diagramme d'activité — circuit déclaration → quittance",
        32,
        width_cm=12.5,
    )
    add_p(
        doc,
        "Le diagramme de déploiement situe les composants : navigateur client, "
        "application Flask/Gunicorn sur Render, base PostgreSQL, service d'email "
        "SendGrid, et environnement de développement local (MySQL).",
    )
    add_diagramme(
        doc,
        "06-deploiement.png",
        "Diagramme de déploiement de BBDA Events",
        33,
        width_cm=13.0,
    )

    add_title(doc, "2.4 Outils et choix techniques", level=2)
    add_p(doc, "Tableau 2 — Stack technique retenue et justifications",
          bold=True, first_line=False, align="center", size=11)
    add_table(
        doc,
        ["Catégorie", "Choix", "Justification résumée"],
        [
            ["Backend", "Python 3.13 + Flask", "Framework léger, MVC explicite, adapté à un périmètre solo"],
            ["Frontend", "HTML5, CSS3, JS vanilla", "Sans framework lourd, conforme au cahier des charges"],
            ["SGBD", "MySQL (dev) / PostgreSQL (prod Render)", "Modèle relationnel fort ; portabilité via SQLAlchemy"],
            ["ORM", "SQLAlchemy", "Sécurité des requêtes, évolution du schéma"],
            ["Auth", "Flask-Login + bcrypt", "Sessions et hachage standards"],
            ["PDF", "ReportLab", "Contrôle fin de la mise en page (quittance)"],
            ["Email", "Flask-Mail / SendGrid (prod)", "Notifications métier"],
            ["Tests", "pytest, Playwright", "Non-régression et parcours bout-en-bout"],
            ["Hébergement", "Render.com", "Déploiement HTTPS public pour démonstration"],
            ["Versioning", "Git / GitHub", "Traçabilité des itérations"],
        ],
    )
    add_p(
        doc,
        "L'architecture suit le patron MVC avec factory pattern (create_app) et "
        "Blueprints par domaine : auth, declarations, agent, admin, public, exports, "
        "notifications, arrieres. La logique métier est isolée dans backend/ ; les "
        "templates et assets dans frontend/.",
    )

    add_title(doc, "2.5 Modélisation des données et règles de gestion", level=2)
    add_p(
        doc,
        "Le schéma relationnel compte une douzaine de tables principales : utilisateur, "
        "organisateur, declaration, liste_artiste, evaluation_agent, paiement, quittance, "
        "arriere, notification, alerte_surveillance, message_contact, parametres_systeme. "
        "Chaque déclaration suit un cycle de statuts explicite.",
    )
    add_p(doc, "Tableau 4 — Cycle des statuts d'une déclaration",
          bold=True, first_line=False, align="center", size=11)
    add_table(
        doc,
        ["Statut", "Signification"],
        [
            ["nouvelle", "Déclaration soumise, non encore prise en charge"],
            ["en_evaluation", "Ouverte par un agent"],
            ["en_attente", "Complément d'information demandé (commentaire obligatoire)"],
            ["montant_fixe", "Tarif et redevance validés"],
            ["paiement_en_attente", "En attente d'encaissement"],
            ["payee", "Solde à zéro après un ou plusieurs versements"],
            ["quittance_delivree", "PDF généré ; éventuelle publication si promotion"],
        ],
    )
    add_p(doc, "Tableau 3 — Principales règles métier (extrait)",
          bold=True, first_line=False, align="center", size=11)
    add_table(
        doc,
        ["Code", "Règle"],
        [
            ["RM-010", "Arriéré ≥ seuil (1 000 FCFA par défaut) → pas de nouvelle déclaration"],
            ["RM-030", "Montant fixé manuellement par l'agent (pas de calcul auto)"],
            ["RM-032", "Montant total = Tarif + Redevance"],
            ["RM-040", "Paiement au guichet BBDA (pas de paiement en ligne dans le prototype)"],
            ["RM-048", "Quittance générée seulement lorsque le solde restant = 0"],
            ["RM-021", "Promotion publique visible seulement si quittance_delivree"],
        ],
    )

    add_page_break(doc)

    # ========== CHAPITRE 3 ==========
    add_title(doc, "CHAPITRE 3 — PRÉSENTATION ET ANALYSE DES RÉSULTATS", level=1)
    add_p(
        doc,
        "Les résultats sont présentés suivant l'ordre des objectifs spécifiques et du "
        "développement itératif. Les captures d'écran proviennent de l'application "
        "réelle (environnement de démonstration). L'URL publique de référence est "
        "https://bbda-events.onrender.com.",
    )

    add_p(doc, "Tableau 5 — Objectifs spécifiques et livrables",
          bold=True, first_line=False, align="center", size=11)
    add_table(
        doc,
        ["Objectif", "Livrable / preuve"],
        [
            ["Analyse processus", "Chap. 1–2 ; fiches / photos quittance"],
            ["Besoins", "Cahier des charges V2 ; REGLES_METIER.md"],
            ["Conception", "ARCHITECTURE.md ; DATABASE_SCHEMA.md"],
            ["Déclaration en ligne", "Module declarations ; Fig. 5–11"],
            ["Espace agent", "Module agent ; Fig. 13–19"],
            ["Notifications email", "email_service ; Fig. 24"],
            ["Arriérés", "moteur.py ; Fig. 6, 14, 15"],
            ["Quittance PDF", "pdf_generator.py ; Fig. 21–23, 28"],
            ["Tests / doc", "~100+ tests pytest ; présent mémoire"],
        ],
    )

    # 3.1 Accueil
    add_title(doc, "3.1 Socle applicatif et page d'accueil publique", level=2)
    add_p(
        doc,
        "Le socle Flask (factory, blueprints, init_db / demo_data) a été mis en place "
        "en premier. La page d'accueil publique présente la marque BBDA Events, le "
        "parcours en trois étapes et les accès inscription / connexion.",
    )
    add_figure(doc, "07-accueil-v2.png", "Page d'accueil publique de la plateforme BBDA Events", 1)

    # 3.2 Auth
    add_title(doc, "3.2 Authentification et contrôle d'accès", level=2)
    add_p(
        doc,
        "L'inscription crée un compte organisateur ; la connexion repose sur Flask-Login "
        "et bcrypt. Un décorateur role_required restreint les routes agent et admin. "
        "Les messages flash informent l'utilisateur en cas d'erreur.",
    )
    add_figure(doc, "02-inscription.png", "Formulaire d'inscription d'un organisateur", 2)
    add_figure(doc, "03-connexion.png", "Formulaire de connexion", 3)
    add_figure(doc, "08-flash-erreur.png", "Exemple de message d'erreur (compte inconnu)", 4)

    # 3.3 Dashboard orga
    add_title(doc, "3.3 Tableau de bord de l'organisateur", level=2)
    add_p(
        doc,
        "Chaque organisateur visualise ses déclarations et leur statut. En cas d'arriéré "
        "bloquant (RM-010), une bannière empêche la création de nouvelles déclarations "
        "jusqu'à régularisation.",
    )
    add_figure(doc, "09-dashboard-orga1.png", "Tableau de bord d'un organisateur avec déclarations en cours", 5)
    add_figure(doc, "10-dashboard-orga4-bloque.png",
               "Tableau de bord d'un organisateur bloqué pour cause d'arriéré", 6)

    # 3.4 Formulaire
    add_title(doc, "3.4 Formulaire de déclaration d'événement", level=2)
    add_p(
        doc,
        "Le formulaire reprend les champs de la fiche papier (identité, nature de la "
        "manifestation, salle, date, artistes) avec comportement dynamique JavaScript "
        "(ajout de lignes d'artistes) et validation client/serveur (RM-011, RM-012).",
    )
    add_figure(doc, "11-formulaire-concert.png", "Formulaire de déclaration d'un concert", 7)
    add_figure(doc, "12-formulaire-festival-artistes.png",
               "Formulaire de déclaration d'un festival avec plusieurs artistes", 8)
    add_figure(doc, "13-formulaire-erreurs.png", "Messages d'erreur lors d'une saisie incomplète", 9)

    # 3.5 Détail
    add_title(doc, "3.5 Page de détail et chronologie", level=2)
    add_p(
        doc,
        "La page de détail présente l'état complet du dossier et une chronologie visuelle "
        "(soumission, montant fixé, paiement, quittance).",
    )
    add_figure(doc, "15-detail-orga3-0.png", "Page de détail d'une déclaration — vue d'ensemble", 10)
    add_figure(doc, "15-detail-orga3-1.png", "Chronologie du traitement d'une déclaration", 11)
    add_figure(doc, "16-detail-montant-fixe.png", "Détail après fixation du montant par l'agent", 12)

    # 3.6 Agent
    add_title(doc, "3.6 Espace agent — pilotage et traitement", level=2)
    add_p(
        doc,
        "L'agent dispose d'indicateurs agrégés, de listes filtrables, d'une vue "
        "surveillance et d'une vue arriérés. L'ouverture d'une déclaration bascule son "
        "statut en « en_evaluation ». L'agent saisit Tarif et Redevance ; le total est "
        "calculé en temps réel côté client (RM-030 à RM-035).",
    )
    add_figure(doc, "17-dashboard-agent.png", "Tableau de bord de l'agent BBDA", 13)
    add_figure(doc, "18-agent-surveillance.png", "Liste des comptes organisateurs sous surveillance", 14)
    add_figure(doc, "19-agent-arrieres.png", "Liste des organisateurs en situation d'arriéré", 15)
    add_figure(doc, "20-agent-traitement.png", "Page de traitement d'une déclaration par l'agent", 16)
    add_figure(doc, "21-avant-validation-total.png",
               "Calcul en temps réel du montant total (tarif + redevance)", 17)

    # 3.7 Paiement
    add_title(doc, "3.7 Confirmation du paiement", level=2)
    add_p(
        doc,
        "Conformément à RM-040, le paiement s'effectue au guichet ; l'agent enregistre "
        "ensuite le mode (espèces, chèque, Orange Money) et le caractère intégral ou "
        "partiel. Un paiement partiel crée automatiquement un arriéré avec échéance à "
        "sept jours (RM-044). Plusieurs versements successifs sont possibles jusqu'à "
        "solde zéro (RM-047, RM-048).",
    )
    add_figure(doc, "24-formulaire-paiement-vide.png", "Formulaire de confirmation de paiement (vierge)", 18)
    add_figure(doc, "25-formulaire-paiement-rempli.png",
               "Formulaire de confirmation de paiement rempli (exemple)", 19)
    add_figure(doc, "27-dossier-apres-paiement.png",
               "Détail de la déclaration après paiement (quittance délivrée)", 20)

    # 3.8 Quittance
    add_title(doc, "3.8 Génération automatique de la quittance PDF", level=2)
    add_p(
        doc,
        "Dès que le solde restant est nul, une quittance PDF est générée automatiquement "
        "via ReportLab (RM-050 à RM-054). La mise en page reproduit le formulaire "
        "physique du BBDA (en-tête, numéro séquentiel, tableau des droits, montant en "
        "lettres, signature). Le téléchargement est réservé au propriétaire du dossier.",
    )
    add_figure(doc, "30-quittance-pdf-v3.png", "Quittance PDF générée automatiquement", 21)
    add_figure(doc, "31-quittance-partiel.png",
               "Document illustrant le cas d'un paiement partiel / reste à payer", 22)
    add_figure(doc, "32-detail-avec-bouton-telecharger.png",
               "Bouton de téléchargement de la quittance (espace organisateur)", 23)

    # Photo quittance réelle
    jpg_files = sorted(IMAGES.glob("*.jpg"))
    if jpg_files:
        add_p(
            doc,
            "La Figure 28 présente une photographie de la quittance papier utilisée "
            "comme référence de conception graphique.",
        )
        add_photo(doc, jpg_files[0],
                  "Quittance papier BBDA (référence terrain photographiée pendant le stage)", 28,
                  width_cm=11.0)

    # Extrait de code
    add_p(doc, "Extrait de code — principe de génération PDF (backend/exports/pdf_generator.py) :",
          first_line=False, italic=True)
    add_code_block(
        doc,
        "def _en_tete(c, quittance):\n"
        "    \"\"\"EN-TETE 3 colonnes : coordonnees BBDA, logo, encadre 'QUITTANCE'.\"\"\"\n"
        "    c.setFont(\"Helvetica-Bold\", 10)\n"
        "    c.drawString(_x(MARGE), _y(15), \"BUREAU BURKINABÈ\")\n"
        "    c.drawString(_x(MARGE), _y(20), \"DU DROIT D'AUTEUR\")\n"
        "    # ... logo + cadre numéro de quittance ...",
    )

    # 3.9 Notifications
    add_title(doc, "3.9 Notifications automatiques par courrier électronique", level=2)
    add_p(
        doc,
        "À chaque étape clé (soumission, montant fixé, quittance disponible, rappel "
        "d'arriéré, alerte surveillance), un courriel est consigné en base puis envoyé "
        "(SMTP local en développement, SendGrid en production Render). Un échec d'envoi "
        "n'interrompt jamais le traitement métier.",
    )
    add_figure(doc, "33-email-montant-fixe.png",
               "Email automatique envoyé lorsque le montant est fixé", 24)

    # 3.10 Arrieres
    add_title(doc, "3.10 Moteur et interface de gestion des arriérés", level=2)
    add_p(
        doc,
        "Le module backend/arrieres/moteur.py centralise : calcul de l'état d'arriéré, "
        "seuil bloquant paramétrable (1 000 FCFA par défaut), échéance à sept jours, "
        "rappels espacés d'au moins sept jours, blocage / déblocage, mise sous "
        "surveillance et report des arriérés dans le calcul d'exigibilité. L'interface "
        "agent permet d'agir sans passer par la base de données.",
    )

    # 3.11 Admin
    add_title(doc, "3.11 Espace administrateur et statistiques", level=2)
    add_p(
        doc,
        "L'administrateur dispose d'un tableau de bord, de la gestion des utilisateurs "
        "(création d'agents, activation / désactivation), des paramètres système "
        "(seuil d'arriéré, délai de rappel) et d'une page de statistiques (déclarations "
        "et redevances mensuelles, répartition par type, classement organisateurs).",
    )
    add_figure(doc, "06-espace-admin.png", "Espace administrateur BBDA Events", 27)

    # 3.12 Promotion
    add_title(doc, "3.12 Face publique et module de promotion", level=2)
    add_p(
        doc,
        "Sans connexion, le public accède à l'accueil, au listing des événements "
        "volontairement promus et déjà quittancés (RM-021), au support (FAQ), au "
        "contact et aux pages légales. L'organisateur peut joindre une affiche "
        "(JPG/PNG, 2 Mo max) et autoriser l'affichage de ses coordonnées.",
    )
    add_figure(doc, "35-formulaire-promotion.png",
               "Section « Promotion publique » du formulaire de déclaration", 26)
    add_figure(doc, "34-detail-evenement-public.png",
               "Page de détail d'un événement sur la face publique", 25)

    # 3.13 Tests
    add_title(doc, "3.13 Tests et validation", level=2)
    add_p(
        doc,
        "La validation s'appuie sur une suite pytest croissante au fil des prompts "
        "(de l'ordre de 50 tests après la quittance à plus de 100 après la face "
        "publique et les tests transverses). Un parcours Playwright vérifie "
        "l'enchaînement bout-en-bout : inscription → déclaration → traitement agent → "
        "paiement → téléchargement de quittance. Un script demo_data.py recharge un "
        "jeu de données pour la soutenance.",
    )

    # Discussion
    add_title(doc, "3.14 Analyse des résultats / Discussion", level=2)
    add_p(
        doc,
        "Au regard de l'objectif général, la plateforme couvre l'intégralité du circuit "
        "cible : déclaration, évaluation, paiement guichet, quittance, notifications, "
        "arriérés et promotion conditionnelle. Par rapport au processus papier, les "
        "gains attendus portent sur : (i) la suppression des déplacements pour la seule "
        "phase de déclaration ; (ii) la traçabilité des statuts et de l'historique ; "
        "(iii) l'automatisation des documents et des alertes ; (iv) le pilotage des "
        "impayés. Ces gains sont démontrés qualitativement par l'artefact et les tests ; "
        "une mesure chronométrée avant/après sur un échantillon réel de dossiers DEPAJ "
        "n'a pas encore été formalisée. [À COMPLÉTER SI POSSIBLE avec 5–10 dossiers "
        "anonymisés : délai moyen papier vs plateforme.]",
    )

    add_title(doc, "3.15 Critiques et suggestions", level=2)
    add_p(doc, "Limites du prototype :", first_line=False)
    add_p(doc, "— Absence de paiement en ligne (choix volontaire RM-040, aligné sur le guichet BBDA).", first_line=False)
    add_p(doc, "— Billetterie annoncée hors périmètre (page placeholder).", first_line=False)
    add_p(doc, "— Stockage des PDF sur disque serveur : sur Render, régénération à la demande "
          "si le fichier éphémère disparaît.", first_line=False)
    add_p(doc, "— WhatsApp non implémenté (API payante).", first_line=False)
    add_p(
        doc,
        "Suggestions : intégrer un moyen de paiement mobile money en ligne ; coupler "
        "BBDA Events à B-GEOLOC / à la plateforme œuvres ; mener un pilote DEPAJ sur "
        "une saison culturelle avec indicateurs de délai et de taux de recouvrement.",
    )

    add_title(doc, "3.16 Bilan du stage", level=2)
    add_p(
        doc,
        "Le stage a permis de conduire un projet informatique complet : du besoin métier "
        "à un produit déployé, testé et documenté. Les compétences mobilisées couvrent "
        "l'analyse, la modélisation relationnelle, le développement full-stack léger, "
        "les tests et le déploiement cloud. [À COMPLÉTER : ressenti personnel, "
        "difficultés rencontrées, apports professionnels concrets.]",
    )

    add_page_break(doc)

    # ========== CONCLUSION ==========
    add_title(doc, "CONCLUSION GÉNÉRALE", level=0)
    add_p(
        doc,
        "L'objectif général de ce mémoire était de concevoir et développer une plateforme "
        "web permettant au BBDA de gérer les déclarations d'événements culturels "
        "occasionnels de la saisie organisateur jusqu'à la quittance, avec suivi des "
        "paiements, arriérés et notifications.",
    )
    add_p(
        doc,
        "Principaux résultats. Tous les objectifs spécifiques opérationnels ont été "
        "atteints : analyse du processus, formalisation des règles, architecture MVC "
        "Flask, modules déclaration / agent / admin / public, emails, moteur d'arriérés, "
        "quittance PDF, tests et déploiement. La plateforme BBDA Events est accessible "
        "publiquement pour démonstration.",
    )
    add_p(
        doc,
        "Confrontation à H1. La dématérialisation implémentée réduit le nombre d'étapes "
        "manuelles (plus de fiche papier à ressaisir, quittance et notifications "
        "automatiques, tableaux de bord). En l'absence d'une mesure chronométrée "
        "officielle avant/après sur dossiers réels, H1 est considérée comme partiellement "
        "confirmée sur le plan fonctionnel et qualitative, et reste à quantifier "
        "empiriquement avec la DEPAJ.",
    )
    add_p(
        doc,
        "Confrontation à H2. Les mécanismes de seuil bloquant, de rappels, de surveillance "
        "et de report d'arriérés sur les dossiers suivants sont effectivement en place et "
        "testés. H2 est donc confirmée sur le plan des dispositifs techniques de "
        "recouvrement ; l'impact financier réel (baisse des pertes) dépendra de "
        "l'appropriation par les agents et du volume traité en production.",
    )
    add_p(
        doc,
        "Apports. Sur le plan pratique, le BBDA dispose d'un prototype opérationnel "
        "couvrant un processus jusque-là non numérisé. Sur le plan académique, le travail "
        "illustre une démarche de recherche appliquée en informatique : problématisation, "
        "hypothèses, conception, réalisation, tests et discussion des limites.",
    )
    add_p(
        doc,
        "Limites. Limites théoriques (revue académique encore perfectible), empiriques "
        "(peu de mesures quantitatives de performance) et méthodologiques (UML formels "
        "incomplets, échantillon d'observation limité au service d'accueil).",
    )
    add_p(
        doc,
        "Perspectives. Paiement en ligne sécurisé ; notifications WhatsApp ; billetterie ; "
        "interopérabilité avec les outils 2025 du BBDA ; industrialisation (sauvegardes, "
        "supervision) ; étude d'impact après mise en production pilote.",
    )

    add_page_break(doc)

    # ========== RÉFÉRENCES ==========
    add_title(doc, "RÉFÉRENCES BIBLIOGRAPHIQUES", level=0)
    add_p(doc, "Les références sont regroupées en ouvrages / documents, articles et sites "
          "internet, par ordre alphabétique du premier auteur ou de l'organisme.",
          italic=True, first_line=False)

    add_title(doc, "Ouvrages, mémoires, rapports et documents institutionnels", level=2)
    refs_ouv = [
        "Bureau Burkinabè du Droit d'Auteur (2025-2026). Protocole de stage — Conception et développement d'une plateforme web de gestion des déclarations… Mémoire de fin d'études en Licence. BBDA / Université Aube Nouvelle, Ouagadougou.",
        "Bureau Burkinabè du Droit d'Auteur (s.d.). Cahier des charges BBDA Events, V2. Document interne, Ouagadougou.",
        "Bureau Burkinabè du Droit d'Auteur (s.d.). Guide complet de développement BBDA Events, V2. Document interne, Ouagadougou.",
        "Bureau Burkinabè du Droit d'Auteur / OMPI (2011). Guide pratique pour les utilisateurs d'œuvres protégées (éd. révisée). WIPO Lex.",
        "Burkina Faso (2000). Arrêté n° 01-053/MAC/SG/BBDA du 20 mars 2000 relatif aux règles de perception. WIPO Lex.",
        "Burkina Faso (2019). Loi n° 048-2019/AN du 12 novembre 2019 portant protection de la propriété littéraire et artistique. WIPO Lex.",
        "Université Aube Nouvelle (2026). Initiation à la méthodologie de recherche — Cours de rédaction scientifique. Licence 3, Tronc commun, janvier 2026.",
    ]
    for r in refs_ouv:
        add_p(doc, r, first_line=False, space_after=6)

    add_title(doc, "Articles de presse et contributions", level=2)
    refs_art = [
        "Burkina24 (2022, 30 mai). BBDA : une cagnotte de plus de 700 millions à répartir à partir du 31 mai 2022. Burkina24.",
        "Digital Magazine Burkina (2025, 28 avril). Le BBDA digitalise la gestion des droits d'auteur… avec B-GEOLOC et sa plateforme numérique. Digital Magazine Burkina.",
        "Kulture Kibaré (2020, 20 juillet). Propriété littéraire et artistique : la nouvelle loi 048 et ses innovations intéressantes. Kulture Kibaré.",
        "Kulture Kibaré (2022, 15 juin). Droits d'auteur : comprendre le fonctionnement des séances occasionnelles. Kulture Kibaré.",
        "leFaso.net (2025, 24 avril). Le BBDA lance officiellement sa plateforme numérique d'adhésion et de déclaration d'œuvres. leFaso.net.",
        "Sawadogo P. F. (2020, 6 novembre). BBDA : Le processus de collecte des droits de la séance occasionnelle. Infos Culture du Faso.",
        "Sidwaya (2024). Droits d'auteur au Burkina : une cagnotte de 397 885 753 FCFA pour 7 307 bénéficiaires. Sidwaya.info.",
    ]
    for r in refs_art:
        add_p(doc, r, first_line=False, space_after=6)

    add_title(doc, "Sites internet", level=2)
    refs_web = [
        "https://bbda.bf/ (site officiel du BBDA) — consulté en juillet 2026.",
        "https://bbda-events.onrender.com (plateforme BBDA Events développée dans le cadre du stage) — consulté en juillet 2026.",
        "https://www.wipo.int/wipolex/fr/legislation/details/19510 (Loi 048-2019/AN) — consulté en juillet 2026.",
        "https://www.wipo.int/wipolex/en/legislation/details/6362 (Arrêté 01-053/2000) — consulté en juillet 2026.",
        "https://sodav.sn/utilisateur/ (SODAV — utilisateurs / séances occasionnelles) — consulté en juillet 2026.",
        "https://clients.sacem.fr/ (SACEM — autorisations spectacles) — consulté en juillet 2026.",
        "https://www.samro.org.za/user (SAMRO — music users) — consulté en juillet 2026.",
        "https://fr.wikipedia.org/wiki/Bureau_burkinabè_des_droits_d'auteurs — consulté en juillet 2026. [À remplacer de préférence par une source primaire BBDA.]",
    ]
    for r in refs_web:
        add_p(doc, r, first_line=False, space_after=6)

    add_page_break(doc)

    # ========== ANNEXES ==========
    add_title(doc, "ANNEXES", level=0)
    add_p(doc, "Plan des annexes (pagination romaine à appliquer en mise en forme finale) :",
          first_line=False)
    add_p(doc, "Annexe I — Photographies de la quittance papier BBDA (référence)", first_line=False)
    add_p(doc, "Annexe II — Captures d'écran complémentaires", first_line=False)
    add_p(doc, "Annexe III — Extraits de code significatifs", first_line=False)
    add_p(doc, "Annexe IV — Schéma de base de données", first_line=False)
    add_p(doc, "Annexe V — Liste des règles métier (renvoi au fichier REGLES_METIER.md)", first_line=False)
    add_p(doc, "Annexe VI — Diagrammes UML (séquence et reprise)", first_line=False)
    add_p(doc, "Annexe VII — [À COMPLÉTER] Scans fiches papier déclaration / évaluation", first_line=False)

    add_title(doc, "Annexe I — Photographies de la quittance papier", level=2)
    for i, jp in enumerate(jpg_files[:3]):
        add_photo(doc, jp, f"Quittance / document papier BBDA — cliché {i+1}", 28 + i, width_cm=10.5)

    add_title(doc, "Annexe II — Captures complémentaires", level=2)
    for fname, leg in [
        ("04-espace-organisateur.png", "Espace organisateur (vue alternative)"),
        ("05-espace-agent.png", "Espace agent (vue alternative)"),
        ("14-apres-soumission.png", "Confirmation après soumission"),
        ("22-dashboard-agent-apres.png", "Dashboard agent après fixation du montant"),
        ("26-dashboard-agent-apres-paiement.png", "Dashboard agent après paiement"),
    ]:
        if (SHOTS / fname).exists():
            add_figure(doc, fname, leg, 100)  # num indicatif annexe

    add_title(doc, "Annexe III — Extrait de code (quittance)", level=2)
    code_path = Path(r"C:\bbda_events\backend\exports\pdf_generator.py")
    if code_path.exists():
        lines = code_path.read_text(encoding="utf-8", errors="replace").splitlines()
        excerpt = "\n".join(lines[72:110])
        add_code_block(doc, excerpt, "Extrait de pdf_generator.py (en-tête de quittance)")

    add_title(doc, "Annexe IV — Schéma de base de données", level=2)
    add_p(
        doc,
        "Le fichier memoire/database/schema_diagram.pdf contient le diagramme "
        "relationnel. Le script SQL schema.sql décrit les tables. "
        "Le détail narratif et technique figure dans DATABASE_SCHEMA.md. "
        "[Insérer manuellement le PDF schema_diagram.pdf si besoin : Insertion > Objet.]",
        first_line=False,
        italic=True,
    )

    add_title(doc, "Annexe VI — Diagrammes UML (séquence et reprise)", level=2)
    add_p(
        doc,
        "Les diagrammes de séquence détaillent deux parcours critiques. L'ensemble "
        "des sources (.mmd) et images (.png) est disponible dans "
        "memoire/docs/diagrammes/.",
        first_line=False,
    )
    add_diagramme(
        doc,
        "03-sequence-declaration.png",
        "Diagramme de séquence — inscription et déclaration d'événement",
        34,
        width_cm=13.5,
    )
    add_diagramme(
        doc,
        "04-sequence-paiement-quittance.png",
        "Diagramme de séquence — évaluation, paiement et quittance",
        35,
        width_cm=13.5,
    )
    add_p(
        doc,
        "Pour la soutenance, les six diagrammes du dossier diagrammes/ peuvent être "
        "projetés séparément (cas d'utilisation, classes, deux séquences, activité, "
        "déploiement).",
        first_line=False,
        italic=True,
    )

    add_page_break(doc)

    # ========== POINTS À COMPLÉTER ==========
    add_title(doc, "ANNEXE FINALE — POINTS À COMPLÉTER / VÉRIFIER / CORRIGER", level=0)
    add_p(
        doc,
        "Cette section liste explicitement ce que l'étudiant doit compléter, vérifier "
        "auprès du BBDA / de l'université, ou remplacer par une information personnelle. "
        "Les passages équivalents déjà rédigés dans le corps du texte sont des "
        "propositions de même nature (à valider).",
        first_line=False,
        italic=True,
    )
    items = [
        "Page de garde : coller le modèle officiel U-AUBEN si différent ; logo établissement.",
        "Directeur de mémoire : nom, grade, signature.",
        "Maître de stage : nom, fonction exacte à la DEPAJ / BBDA.",
        "Dédicace : personnaliser le texte proposé.",
        "Remerciements : remplacer les [À COMPLÉTER] par les vrais noms.",
        "Dates exactes du stage (jour/mois/année début–fin).",
        "Date et lieu de soutenance.",
        "Vérifier avec le maître de stage le vocabulaire officiel : « déclaration » vs « autorisation » / délai 72 h (presse).",
        "Obtenir et scanner : fiche de déclaration papier + fiche d'évaluation (Annexes VII).",
        "Insérer schema_diagram.pdf dans l'Annexe IV (si le jury veut le PDF ER en plus des UML).",
        "Relire les diagrammes UML (memoire/docs/diagrammes/) avec le maître de stage / encadrant.",
        "Mesure empirique H1 (optionnel mais valorisant) : délai moyen de traitement papier vs plateforme sur un petit échantillon anonymisé.",
        "Confirmer validation formelle du cahier des charges par le maître de stage.",
        "Remplacer / compléter la capture 33-email si une capture mail client plus nette est disponible.",
        "Revérifier les PDF officiels Loi 048 et Arrêté 2000 sur WIPO Lex avant citation orale.",
        "Ajouter 2–3 références académiques (CISAC, OMPI, articles sur e-gouvernement) si le jury l'exige.",
        "Pagination : chiffres romains jusqu'au sommaire, arabes à partir de l'Introduction (à régler dans Word : sauts de section).",
        "Table des matières automatique Word (Références > Table des matières) après styles Titre 1/2/3.",
        "Relire orthographe / accords ; harmoniser « séance occasionnelle » / « événement occasionnel ».",
        "Compte-rendu de stage / attestation de stage à joindre si l'école le demande (hors présent document).",
    ]
    for i, it in enumerate(items, 1):
        add_p(doc, f"{i}. {it}", first_line=False, space_after=4)

    add_p(doc, "", first_line=False)
    add_p(
        doc,
        "Document généré automatiquement à partir des sources du dépôt BBDA Events "
        "(protocole, cahier des charges, redaction/*.md, docs/, screenshots/, images/, "
        "RECHERCHE_DOCUMENTAIRE_THEME_BBDA_EVENTS.md) et du Cours_Rédaction scientifique.pptx "
        "(Université Aube Nouvelle, janvier 2026).",
        size=10,
        italic=True,
        first_line=False,
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    print(f"OK -> {OUT}")
    print(f"Taille : {OUT.stat().st_size / 1024 / 1024:.2f} Mo")


if __name__ == "__main__":
    build()
